import re
import shutil
from os import path, listdir, makedirs, chdir

import win32com.client as win32
from docx import Document
from pandas import DataFrame as DF
from win32com.client import constants

state_mapping = dict({
    'mumbai': 'maharashtra',
    "bangalore": "Karnataka",
    'tamil nadu': 'tamilnadu',
    'hyderabad': 'andhra pradesh'
})

billing_location_mapping = dict(
    mumbai='maharashtra',
    andheri='maharashtra',
    kalamboli='maharashtra',
    maharashtra='maharashtra',
    bangalore='karnataka',
    chennai='tamil nadu',
    hyderabad='andhra pradesh',
)
billing_location_mapping.setdefault('')
state_mapping.setdefault(None, '')


class OPF:
    def __init__(self, path):
        # path is the path or the filename of the file.
        self.path = path

        # document is set to None and not to a document object because,
        # we don't know the extension of the file which is saved.
        # if that file is a doc file, it should be converted to docx and then it's document object should be stored.
        self.document = None
        # <editor-fold desc="Getting file name from path">
        # following code can be easily replaced by one of the os functions related to path.
        if '/' in self.path:
            self.file_name = self.path.split('/')[-1]
        else:
            self.file_name = self.path.split('\\')[-1]
        # </editor-fold>

    def seperate_doc(self, docx_folder_name='Docxs', docs_folder_name='Docs'):
        # <editor-fold desc="Converting a doc file to docx file">
        if '.docx' not in self.path and '.doc' in self.path:
            # Control is here means that, file is a doc file but not a docx file.

            # saving current location of the doc file so that,
            # it can be used to move it in another folder once this file is converted to docx file.
            prev_path = self.path

            # saving doc as doc and updating current file location to new file location of docx file.
            self.path = self.save_as_docx()

            # Moving the previous doc file to docs folder in order to remove all clutter.
            self.move_file(file_path=prev_path, to_path=docs_folder_name)
        # </editor-fold>
        # moving the file to docx file to docx folder.
        self.move_file(to_path=docx_folder_name)

    def save_as_docx(self):
        # Code credits: https://stackoverflow.com/questions/38468442/multiple-doc-to-docx-file-conversion-using-python
        # <editor-fold desc="Opening Microsoft word application.">
        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(self.path)
        doc.Activate()
        # </editor-fold>
        # <editor-fold desc="Rename path with .docx">
        new_file_abs = path.abspath(path)
        new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)
        # </editor-fold>
        # <editor-fold desc="Save and Close">
        word.ActiveDocument.SaveAs(
            new_file_abs, FileFormat=constants.wdFormatXMLDocument
        )
        doc.Close(False)
        # </editor-fold>
        # Returning the absoulute path of the docx file generated.
        return new_file_abs

    def move_file(self, to_path, file_path=None):
        # <editor-fold desc="Make Directory if not exists.">
        if file_path is None:
            # File path is none means that, user want to move file whose location is stored in the object instance.
            file_path = self.path
        if not path.exists(to_path):
            # using makedirs instead of makedir to enable making of nested directories.
            # That is, enabling making of child directory even if parent directory is not present at the moment
            # in which child directory was being created by making all the directories recursively.
            makedirs(to_path)
        # </editor-fold>
        # moving the file stored with path given by file_path to path given by to_path.
        shutil.move(file_path, to_path)

    def get_tables(self):
        # <editor-fold desc="setting document object if not defined.">
        if self.document is None:
            # Document object is not modified since the init of the object.

            # setting object's document as Document object of current file's path.
            self.document = Document(self.path)
        # </editor-fold>
        # Now, document object is already defined.
        # returning the list of tables.
        return self.document.tables

    @staticmethod
    def create_table(docx_table):
        table_primitive = []
        # <editor-fold desc="Iterate over all rows and cells and append it to table_primitive list.">
        for row in docx_table.rows:
            # new instance of list for all all the rows in table.
            table_row = []
            for cell in row.cells:
                table_row.append(cell.text.strip('\n').strip(' ').replace('\n', ' '))
            table_primitive.append(table_row)
        # </editor-fold>
        return table_primitive

    @staticmethod
    def print_table(table_data):
        # <editor-fold desc="getting max width of all the columns by transposing and mapping to find the lengths.">
        widths = [max(map(len, col)) for col in zip(*table_data)]
        # </editor-fold>
        # <editor-fold desc="printign the rows by using ljust to make every element of same width.">
        for row in table_data:
            print(" # ".join((val.ljust(width) for val, width in zip(row, widths))))
        # </editor-fold>

    def extract_data(self) -> dict:
        # <editor-fold desc="1. Creating seperate objects for tables used for address and sales.">
        tables = self.get_tables()

        if not tables:
            return dict()
        elif len(tables) == 1:
            raise ValueError('Found only one table, cannot parse it further.')

        address_table = self.create_table(tables[0])
        sales_table = self.create_table(tables[1])
        # </editor-fold>
        # <editor-fold desc="2. Extracting fields from tables and loose fields.">
        address_data = self.parse_address_tables(address_table)
        sales_data = self.parse_sales_data(sales_table)
        loose_fields = self.get_loose_fields()
        # </editor-fold>
        # <editor-fold desc="3. Combining sales, address, loose details to final result.">
        final_result = address_data.copy()
        final_result.update(sales_data)
        final_result.update(loose_fields)
        # </editor-fold>
        # <editor-fold desc="4. Setting the type of gst to final result based on algorithm demonstrated by Anoop Sir.">
        # <editor-fold desc="setting galaxy and buyer's address">
        cbs = (final_result['badstate'] if final_result['badstate'] else '').lower()
        gbs = billing_location_mapping.get(
            (final_result['opf_location'] if final_result['opf_location']
             else ''  # to prevent performing operations on NoneType
             ).replace('/', '').lower()  # one case in which / was present in the field.
        )
        # </editor-fold>
        # <editor-fold desc="Algorithmic switch case.">
        if gbs == cbs:
            type_gst = 'same state'
        elif 'sez' in cbs:
            type_gst = 'sez'
            # if sez is present in the billing address, over writing the value of billing address state to
            # state without parentheses.
            final_result['badstate'] = cbs.split("(")[0]
        else:
            type_gst = 'interstate'
        # </editor-fold>
        # Setting type of gst to the final result.
        final_result['type_gst'] = type_gst
        # </editor-fold>
        # <editor-fold desc="5. Setting value of dc state as gbs.">
        # delivery challan is state corresponding to galaxy billing location. Which is gbs.
        final_result['dc_state'] = gbs
        # </editor-fold>
        return final_result

    @staticmethod
    def position_insensitive_strip(string, weeds):
        # in order to remove an duplicates
        weeds = set(weeds)

        def lstrip(string, weeds):
            while string and string[0] in weeds:
                string = string[1:]
            return string

        all_stripped = lstrip(lstrip(string, weeds)[::-1], weeds)[::-1]
        return all_stripped

    def get_element_from_block(self, block: list, identifier: str, split_by: str) -> str:
        # <editor-fold desc="1. Escaping all common regex literals.">
        identifier = identifier.replace('(', '\(').replace(')', '\)').replace(" ", '\s*') + '\s*'
        # </editor-fold>
        for string in block:
            # searching presence of identifier in string.
            if re.search(identifier, string, flags=re.IGNORECASE):
                # if identifier is present in string, get split the string by split by character and store it .
                probable_result = split_by.join(string.split(split_by)[1:])

                # After splitting if the result is an empty string or string with only spaces means that,
                # splitby character is not present and
                # an other try is given to check if colon was present in the string.
                if not probable_result.strip():
                    probable_result = ":".join(string.split(':')[1:])

                # Stripping the result to remove all the excessive characters from the ends of the result.
                # the weed characters may occur in any order defined by TCS regex given as follows:
                # ('-'+':'+'_'+' ')*
                return self.position_insensitive_strip(probable_result, '-+: ')

    def parse_address_tables(self, address_table):
        # This function
        # 1) takes as an input address table,
        # 2) splits it into billing and delivery.
        # 3) Extract information from both blocks,
        # 4) Return the merged information from both the block.

        def parse_address_block(block):
            which_address = block.pop(0)
            fields = ['State:', 'Contact Person:', "Tel#", "Email#", "GSTN NO:"]
            a = block[0]

            # first_field_index is the index in the block before which address is present.
            first_field_index = None

            # <editor-fold desc="1. Finding the index of row in which first field identifier is found.">
            for index, element in enumerate(block):
                for field in fields:
                    # Fields are stored as field+split_by character.
                    if field[:-1] in element:
                        # field is found in the element, setting the field index and breaking the loop.
                        first_field_index = index
                        break
                if first_field_index is not None:
                    # Sadly, one cannot do multiple breaks to break nested loops,
                    # Hence checking if first_field_index was set in any of the inner loops and break this outer loop.
                    break
            # </editor-fold>
            if first_field_index is None:
                return {}  # an empty dict suggesting failure in parsing all fields.
            # <editor-fold desc="2. setting name and address variables..">
            name = block[0]
            if first_field_index == 1:
                address = "\n".join(block[:first_field_index])
            else:
                address = "\n".join(block[1:first_field_index])
            # </editor-fold>
            block = block[first_field_index:]
            # <editor-fold desc="3. Setting state variable.">
            state = self.get_element_from_block(block, 'State', ':')
            if state is not None and state is not '':
                # state is set with some not empty list of characters.
                if re.search('Mumbai', state, flags=re.IGNORECASE):
                    # Mumbai was a state in some opfs.
                    state = 'Maharashtra'
            if not state:
                # if state is not mentioned and address mentions mumbai,
                # setting maharashtra as a state.
                if re.search('Mumbai', address, flags=re.IGNORECASE):
                    state = 'Maharashtra'
                else:
                    # if state is not found and it is also address not contains mumbai.
                    # setting state to empty string to prevent errors in further processing and computations.
                    state = ''
            if state:
                # mapping wrongly entered cities as states to their corresponding state.

                # Getting corresponding city to state mapping
                res_state = state_mapping.get(state.lower())

                # Setting state as res_state if it returned not None or not empty string
                # else: setting state same as previous state.
                state = res_state if res_state else state
            # </editor-fold>
            # <editor-fold desc="4. Setting contact person, tel and email fields.">
            contact_person = self.get_element_from_block(block, 'Contact Person', ':')
            email = self.get_element_from_block(block, 'Email', '#')
            tel = self.get_element_from_block(block, 'tel', '#')
            # </editor-fold>
            # <editor-fold desc="5. Handling combination of gst and pan no and setting their corresponding values.">
            if self.get_element_from_block(block, 'GST', ':') is not None:
                # gst was found in one of the rest block elements.
                # <editor-fold desc="Finding string in blocks which has 'gst' in it and storing it in variable i.">
                for i in block:
                    if 'GST' in i:
                        break
                # </editor-fold>
                # <editor-fold desc="seperating gst and pan no fields.">
                if 'pan no' in i.lower():
                    # if pan is found entangled with gst field,
                    # splitting string having from pan no and storing parts as gst followed by pan no.
                    gstn, pan = re.split('PAN NO', i, flags=re.IGNORECASE)
                else:
                    # else means that,
                    # gst pan no is not found in the field having gst, hence setting pan number as empty string.
                    gstn, pan = i, ''
                # </editor-fold>
                # <editor-fold desc="splitting both fields with their respective seperators.">
                gstn = gstn.split(':')[-1]
                pan = pan.split(":-")[-1]
                # </editor-fold>

            else:
                # no gst was found in the strings of block.
                gstn = pan = ''
            # </editor-fold>

            # combining all fields and returning the resultant dict.
            return dict(
                name=name,
                address=address, state=state,
                pan=pan, gstn=gstn,
                contact_person=contact_person,
                email=email, tel=tel
            )

        # Getting transpose of address table to make all columns as rows
        # In order to get seperate list for
        #      b) Delivery address and
        #      a) Billing address.
        address_table = [list(i) for i in zip(*address_table)]

        # adding prefixes to differentiate between fields of billing and delivery address.
        billing_address = {'bad' + k: v for k, v in parse_address_block(address_table[0]).items()}
        delivery_address = {'dad' + k: v for k, v in parse_address_block(address_table[1]).items()}

        # updating billing address dict with items of delivery address in order to
        # merge both of the dictionaries' contents and return the merged one.
        billing_address.update(delivery_address)

        return billing_address

    def parse_sales_data(self, sales_table):
        # header is row having corresponding labels for cells in following table.
        #  Sr, Description, Qty., Unit Price, Total Price
        # above line gives the general format of the header list.
        header = sales_table.pop(0)

        result = {}  # The resultant dict which will store sales data.

        # <editor-fold desc="Getting all products' details and their count">
        # i counts the number of products.
        i = 0

        while True:
            i += 1
            # this is a dangerous hard coding done.
            # I've set the sr no as first element of the flattend table.
            sr_no = sales_table[0][0]

            # checking if sr no is having any digits.
            if not [i for i in sr_no if i.isdigit()]:
                # If there are no digits in the sr no,
                # Then, the list of digits [i for i in sr_no if i.isdigit()] will be empty and then breaking the loop.
                break

            # Setting the fields based on hard coded structure of the fields as described in the start of this function.
            result.update({'desc_' + sr_no: sales_table[0][1]})
            result.update({'qty_' + sr_no: sales_table[0][2]})
            result.update({'unit_price_' + sr_no: sales_table[0][3]})
            result.update({'total_price_' + sr_no: sales_table[0][1]})

            # removing the current row from the table so that program always processes on the first list in nested list.
            sales_table.pop(0)

        # decreasing the count to compensate the extra increment
        # that took place before the execution came out of the while loop.
        i -= 1

        # This is done in accordance to a case where
        # there is only one element and then i was set to 1
        # and then when it was decremented with one, it became zero.
        i = max([i, 1])

        # updating the sales details dict with the number of products.
        result.update(dict(number_of_products=i))
        # </editor-fold>

        # <editor-fold desc="getting gst percentages.">
        for i in sales_table:
            for j in i:
                # iterating over each cell in the remaining cells in the sales_block and searching for gst in fields:
                # for each field having following types of gst, instead of splitting and applying lots of brains to it,
                # I extracted all the numbers from the field and converted it to string to store it.
                if 'CGST' in j:
                    result.update(cgst_percentage="".join([k for k in j if k.isdigit()]))
                elif 'SGST' in j:
                    result.update(sgst_percentage="".join([k for k in j if k.isdigit()]))
                elif 'IGST' in j:
                    result.update(igst_percentage="".join([k for k in j if k.isdigit()]))
        # </editor-fold>
        return result

    def get_loose_fields(self):
        # <editor-fold desc="Getting a new object of document if not exists.">
        try:
            # trying to access self.document to ensure that it exists.
            self.document
        except NameError:
            # If document was not found, a nameError will be generated.
            # and hance setting self.document to Document object of current file path.
            self.document = Document(self.path)
        # </editor-fold>
        paragraphs = self.document.paragraphs

        lines = []

        # <editor-fold desc="Appending all paragraphs which are of type x    y or x\ty to lines list.">
        for paragraph in paragraphs:
            for run in paragraph.runs:
                if '\t' in run.text or ' ' * 4 in run.text:
                    lines.append(paragraph)
                    break
        # </editor-fold>

        # line_text is list of texts of lines which have \t replaced with 4 spaces in order to maintain homogeneity.
        line_texts = [line.text.replace('\t', '    ').strip() for line in lines]
        line_texts = [i for i in line_texts if i]

        # final list of all the fields and its values
        # texts will be list of string in the near future.
        texts = list()

        for text in line_texts:
            # for each text in texts, counting max number of spaces.
            count_space = 1
            while ' ' * count_space in text:
                count_space += 1
            else:
                count_space -= 1

            # count space when becomes one, it is decremented by 1 and it becomes zero.
            # in order to prevent ' '*count_spaces to become an empty seperator,
            # count_space is set to 1 if it becomes 0.
            count_space = max([count_space, 1])

            # Splitting the string based on max number of spaces and appending both the elements to the list.
            texts.extend(text.split(count_space * ' '))

        payment_terms = ''
        # THis is hard coded O(n) solution which traverses over all the elements in the list and
        # finds if payment terms is present in that run and append it's value to the final dict of sales dict.
        for i in paragraphs:
            # Getting all the paragraphs row wise.
            if re.search('PAYMENT TERMS', i.text, flags=re.IGNORECASE):
                # payment terms was found in the text of the paragraph.
                # Getting payment terms splitted by :.
                payment_terms = self.get_element_from_block([i.text], 'PAYMENT TERMS', ":")

        # finding opf location that is useful to determine the state from which opf was made.
        opf_location = self.get_element_from_block(texts, 'Galaxy Billing from (Location)', ":")

        # combining all the extracted elements into a single dict in order to return a packed result.
        result_dict = dict(
            sales_person=self.get_element_from_block(texts, 'Sales Person', ":"),
            pot_id=self.get_element_from_block(texts, 'POT ID', ":"),
            opf_no=self.get_element_from_block(texts, 'OPF No.', "."),
            customer_name=self.get_element_from_block(texts, 'Customer Name', ":"),
            opf_date=self.get_element_from_block(texts, 'OPF Date', ":"),
            opf_location=opf_location,
            purch_order_no=self.get_element_from_block(texts, 'Purchase Order', "."),
            payment_terms=payment_terms
        )
        # result was not immediately consumed in order to debug.
        return result_dict


def clean_all_dict_fields(dictionary: dict, all_weeds: str = ':-:;\\ .,\n') -> dict:
    for key, value in dictionary.items():
        dictionary[key] = OPF.position_insensitive_strip(str(value), weeds=all_weeds)
    return dictionary


if __name__ == '__main__':
    docx_folder_name = 'Docxs'

    # traversing all the files in the current directory.
    for file in listdir():
        if '.doc' in file:
            # if file is a doc file, convert it to docx and seggregate doc and docx.
            OPF(path.abspath(file)).seperate_doc()

    # moving the control to docx folder directory
    chdir(path.abspath(docx_folder_name))

    # final list of all dicts for all files.
    result_dict_list = []

    # opf = OPF('OPF- TK-024.docx').extract_data()

    for i, file_name in enumerate(listdir()):
        if '.docx' in file_name:
            # searching for only docx file.
            print(file_name)

            # creating opf instance for all opf docx.
            opf = OPF(file_name)

            # extract data.
            data_dict = opf.extract_data()

            # appending opf link in order it can be joined with other excel file.
            data_dict.update({'opf link': file_name.split('.')[0]})

            # cleaning the data of returned dict.
            data_dict = clean_all_dict_fields(data_dict, all_weeds=': -/;\n')

            # appending the extracted data to the result dict.
            result_dict_list.append(data_dict)

    # removing result of all the dicts which have badstate defined and is not an empty string.
    result_dict_list = [i for i in result_dict_list if i['badstate']]

    # result_dict_list = sorted(result_dict_list, key=lambda x:x['badstate'] if x['badstate'] is not None else '')
    df = DF(result_dict_list)

    # column headers.
    all_keys = list(df.keys())

    header = [
        all_keys.pop(all_keys.index('dc_state')),
        all_keys.pop(all_keys.index('opf_no')),
        all_keys.pop(all_keys.index('customer_name')),
        all_keys.pop(all_keys.index('purch_order_no')),
        all_keys.pop(all_keys.index('opf_date')),
        all_keys.pop(all_keys.index('payment_terms')),
        all_keys.pop(all_keys.index('dadname')),
        all_keys.pop(all_keys.index('dadaddress')),
        all_keys.pop(all_keys.index('dadstate')),
        all_keys.pop(all_keys.index('dadgstn')),
        all_keys.pop(all_keys.index('badname')),
        all_keys.pop(all_keys.index('badaddress')),
        all_keys.pop(all_keys.index('badstate')),
        all_keys.pop(all_keys.index('badgstn')),
        all_keys.pop(all_keys.index('number_of_products')),
        all_keys.pop(all_keys.index('opf_location')),
        all_keys.pop(all_keys.index('type_gst')),
        all_keys.pop(all_keys.index('igst_percentage')),
        all_keys.pop(all_keys.index('sgst_percentage')),
        all_keys.pop(all_keys.index('cgst_percentage')),
    ]

    desc = sorted([i for i in all_keys if 'desc' in i and all([d.isdigit() for d in i[5:]])], key=lambda x: int(x[5:]))
    qty = sorted([i for i in all_keys if 'qty' in i], key=lambda x: int("".join([i for i in x[4:] if i.isdigit()])))
    unit_price = sorted([i for i in all_keys if 'unit_price' in i],
                        key=lambda x: int("".join([i for i in x[11:] if i.isdigit()])))

    all_keys = [i for i in all_keys if i not in desc + unit_price + qty]
    unpacked_tuples = []
    for desc__qty__unit_price in zip(desc, qty, unit_price):
        unpacked_tuples += desc__qty__unit_price

    df = df[header + unpacked_tuples]
    df.to_excel('../final_output.xlsx')  # saving out of docx folder.
